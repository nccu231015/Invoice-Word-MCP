import os
import json
import re
import binascii
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.table import _Cell
from datetime import datetime

# 全局進度回調函數
_progress_callback = None

def set_progress_callback(callback):
    """
    設置進度回調函數

    參數:
    callback -- 回調函數 (step, message, progress, result)
    """
    global _progress_callback
    _progress_callback = callback

def report_progress(step, message, progress=None, result=None):
    """
    報告處理進度

    參數:
    step -- 處理步驟
    message -- 進度消息
    progress -- 完成百分比 (0-100)
    result -- 處理結果
    """
    global _progress_callback
    if _progress_callback:
        _progress_callback(step, message, progress, result)

def set_cell_border(cell, **kwargs):
    """
    設置單元格邊框
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # 檢查是否有邊框屬性
    for key, value in kwargs.items():
        if key == 'top':
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            top = tcBorders.first_child_found_in("w:top")
            if top is None:
                top = OxmlElement('w:top')
                tcBorders.append(top)
            top.set(qn('w:val'), value)
        if key == 'bottom':
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            bottom = tcBorders.first_child_found_in("w:bottom")
            if bottom is None:
                bottom = OxmlElement('w:bottom')
                tcBorders.append(bottom)
            bottom.set(qn('w:val'), value)
        if key == 'left':
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            left = tcBorders.first_child_found_in("w:left")
            if left is None:
                left = OxmlElement('w:left')
                tcBorders.append(left)
            left.set(qn('w:val'), value)
        if key == 'right':
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            right = tcBorders.first_child_found_in("w:right")
            if right is None:
                right = OxmlElement('w:right')
                tcBorders.append(right)
            right.set(qn('w:val'), value)

def set_cell_shading(cell, fill_color):
    """
    設置單元格背景色
    """
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement

def format_number(value):
    """將數字格式化為整數"""
    try:
        if isinstance(value, (int, float)):
            return str(int(value))
        return value
    except:
        return value

def analyze_template(template_path):
    """詳細分析模板檔案的結構並返回關鍵信息"""
    report_progress('analyzing', '正在分析模板結構', 5)
    
    template_info = {
        "placeholders": set(),
        "tables_info": [],
        "item_table_index": -1
    }
    
    try:
        doc = Document(template_path)
        print(f"分析模板: {template_path}")
        print(f"段落數: {len(doc.paragraphs)}")
        print(f"表格數: {len(doc.tables)}")
        
        # 尋找所有段落中的佔位符
        for i, para in enumerate(doc.paragraphs):
            # 顯示段落的原始內容
            text = para.text
            if "{" in text or "}" in text:
                print(f"段落 {i+1} 內容: '{text}'")
                
            # 使用正則表達式尋找佔位符 - 適用於{fieldName}格式
            matches = re.findall(r'{([^{}]+)}', text)
            if matches:
                for match in matches:
                    clean_match = match.strip()
                    template_info["placeholders"].add(clean_match)
                print(f"段落 {i+1}: 找到佔位符 {matches}")
        
        # 分析所有表格
        for i, table in enumerate(doc.tables):
            table_info = {
                "rows": len(table.rows),
                "columns": len(table.rows[0].cells) if len(table.rows) > 0 else 0,
                "placeholders": set(),
                "is_item_table": False
            }
            
            # 檢查表格中的佔位符
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    for p, paragraph in enumerate(cell.paragraphs):
                        cell_text = paragraph.text
                        
                        # 顯示含有大括號的單元格內容
                        if "{" in cell_text or "}" in cell_text:
                            print(f"表格 {i+1}, 行 {r+1}, 列 {c+1}, 段落 {p+1} 內容: '{cell_text}'")
                        
                        # 使用正則表達式尋找佔位符 - 適用於{fieldName}格式
                        matches = re.findall(r'{([^{}]+)}', cell_text)
                        if matches:
                            for match in matches:
                                clean_match = match.strip()
                                table_info["placeholders"].add(clean_match)
                            print(f"表格 {i+1}, 行 {r+1}, 列 {c+1}: 找到佔位符 {matches}")
            
            # 判斷是否為項目表格 (含有類別、項目、單價、數量、金額等表頭)
            if len(table.rows) > 0:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                header_text = " ".join(headers)
                if any(keyword in header_text for keyword in ["類別", "項目", "單價", "數量", "金額"]):
                    table_info["is_item_table"] = True
                    template_info["item_table_index"] = i
                    print(f"表格 {i+1} 被識別為項目表格, 表頭: {headers}")
            
            template_info["tables_info"].append(table_info)
        
        # 總結發現的佔位符
        if template_info["placeholders"]:
            print("\n模板中的所有佔位符:")
            for p in sorted(template_info["placeholders"]):
                print(f"- {p}")
        
        report_progress('analyzing', '模板分析完成', 10)
        return template_info
    except Exception as e:
        error_message = f"分析模板時發生錯誤: {str(e)}"
        print(error_message)
        report_progress('error', error_message, 0)
        return template_info

def replace_text_with_field_value(paragraph, field_mapping):
    """
    使用欄位映射替換段落中的佔位符
    
    參數:
    paragraph -- Document paragraph 對象
    field_mapping -- 欄位映射字典
    
    返回:
    bool -- 是否進行了替換
    """
    text = paragraph.text
    changed = False
    
    # 不處理 {#items} 和 {/items} 標籤，以及項目表格相關的佔位符
    skip_placeholders = ["{#items}", "{/items}", "{category}", "{items}", "{unit}", "{quantity}", "{amount}"]
    for skip in skip_placeholders:
        if skip in text:
            return False
        
    # 使用正則表達式找出所有 {fieldName} 格式的佔位符
    pattern = r'{([^{}]+)}'
    matches = re.findall(pattern, text)
    
    if not matches:
        return False
    
    # 保存所有run的格式信息
    runs_info = []
    for run in paragraph.runs:
        runs_info.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name,
            'size': run.font.size,
            'color': run.font.color.rgb if run.font.color else None
        })
    
    # 對於每個找到的佔位符，嘗試替換
    new_text = text
    for field in matches:
        clean_field = field.strip()
        if clean_field in field_mapping:
            placeholder = f"{{{clean_field}}}"
            replacement = field_mapping[clean_field]
            new_text = new_text.replace(placeholder, str(replacement))
            print(f"替換: '{placeholder}' → '{replacement}'")
            changed = True
        else:
            # 只有當不是項目表格相關的佔位符時才顯示警告
            if clean_field not in ["category", "items", "unit", "quantity", "amount"]:
                print(f"警告: 佔位符 '{clean_field}' 在欄位對應中不存在")
    
    if changed:
        print(f"替換前: '{text}'")
        print(f"替換後: '{new_text}'")
        
        # 清空段落
        paragraph.clear()
        
        # 如果原來只有一個run，直接使用原始格式
        if len(runs_info) == 1:
            run = paragraph.add_run(new_text)
            info = runs_info[0]
            run.bold = info['bold']
            run.italic = info['italic']
            run.underline = info['underline']
            if info['font_name']:
                run.font.name = info['font_name']
            if info['size']:
                run.font.size = info['size']
            if info['color']:
                run.font.color.rgb = info['color']
        else:
            # 複雜情況，簡單添加文字但嘗試保留部分格式
            # 檢查是否大部分run有相同的基本格式
            has_common_bold = all(run['bold'] == runs_info[0]['bold'] for run in runs_info if run['bold'] is not None)
            has_common_italic = all(run['italic'] == runs_info[0]['italic'] for run in runs_info if run['italic'] is not None)
            has_common_font = all(run['font_name'] == runs_info[0]['font_name'] for run in runs_info if run['font_name'] is not None)
            
            run = paragraph.add_run(new_text)
            if has_common_bold:
                run.bold = runs_info[0]['bold']
            if has_common_italic:
                run.italic = runs_info[0]['italic']
            if has_common_font and runs_info[0]['font_name']:
                run.font.name = runs_info[0]['font_name']
            
        return True
    
    return False

def apply_cell_style(cell, style=None):
    """套用單元格樣式，例如對齊方式和填充"""
    if not style:
        return
    
    # 設置對齊方式
    if style.get("align") == "center":
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif style.get("align") == "right":
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 設置背景顏色
    if style.get("fill_color"):
        set_cell_shading(cell, style.get("fill_color"))
    
    # 設置字體粗體
    if style.get("bold"):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # 設置垂直對齊
    if style.get("vertical_align") == "center":
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def cleanup_temp_files(temp_dir, file_prefix="quote_"):
    """清理暫存檔案，避免權限問題"""
    try:
        for file_name in os.listdir(temp_dir):
            if file_name.startswith(file_prefix):
                try:
                    file_path = os.path.join(temp_dir, file_name)
                    os.remove(file_path)
                    print(f"已刪除舊檔案: {file_path}")
                except:
                    pass
    except:
        pass

def format_date(date_str):
    """
    將日期格式轉換為 YYYY/MM/DD 格式
    """
    try:
        if not date_str:
            return ""
        # 如果是已經符合 YYYY/MM/DD 格式的，直接返回
        if re.match(r'\d{4}/\d{2}/\d{2}', date_str):
            return date_str
        # 嘗試將 YYYY-MM-DD 格式轉換為 YYYY/MM/DD
        if re.match(r'\d{4}-\d{2}-\d{2}', date_str):
            year, month, day = date_str.split("-")
            return f"{year}/{month}/{day}"
        return date_str
    except:
        return date_str

def create_field_mapping(quote):
    """
    根據報價單數據創建佔位符映射
    
    參數:
    quote -- 單個報價單的字典數據
    
    返回:
    dict -- 佔位符到值的映射
    """
    try:
        # 確保 quote 有 header 字段
        if not quote or "header" not in quote:
            raise KeyError("報價單數據缺少 'header' 字段")
            
        header = quote["header"]
        
        # 添加調試日誌
        print(f"處理報價單: {header.get('quoteNumber', 'Unknown')}")
        print(f"Header 內容: {header}")
        print(f"Quote 其他欄位: {[k for k in quote.keys() if k != 'header']}")
        
        # 創建欄位映射 (佔位符名稱 -> 值)
        field_mapping = {
            # 標題及基本資訊
            "title": header.get("Title", ""),
            "quoteNumber": header.get("quoteNumber", ""),
            
            # 客戶資訊
            "clientName": header.get("recipient", ""),
            "clientContact": header.get("companyContact", ""),
            "clientEmail": header.get("companyEmail", ""),
            "quoteDate": format_date(header.get("start_date", "")),
            "validUntil": format_date(header.get("end_date", "")),
            "recipient": header.get("recipient", ""),
            
            # 公司資訊
            "companyName": header.get("companyName", "亦式數位互動有限公司"),
            "companyContact": header.get("companyContact", "0988363357"),
            "companyEmail": header.get("companyEmail", "istudiodesign.tw@gmail.com"),
            "unifiedNumber": header.get("key", "96790278"),
            "staff": header.get("staff", ""),
            "key": header.get("key", "96790278"),
            
            # 總計資訊（動態計算）
            "subtotal": format_number(quote.get('total_without_tax', 0)),
            "discountPercentage": str(int((quote.get('discount', 0) / quote.get('total_without_tax', 1)) * 100)) if quote.get('total_without_tax', 0) > 0 else "0", 
            "discount": format_number(quote.get('discount', 0)),
            "taxRate": str(int((quote.get('tax_rate', 0) / (quote.get('total_without_tax', 1) - quote.get('discount', 0))) * 100)) if (quote.get('total_without_tax', 0) - quote.get('discount', 0)) > 0 else "5", 
            "tax": format_number(quote.get('tax_rate', 0)),
            "total": format_number(quote.get('total_with_tax', 0)),
            
            # 支付詳情和備註
            "paymentDetails": "付款方式：銀行轉賬",
            "notes": quote.get("notes", "新客戶享有9折優惠"),
        }
        
        return field_mapping
    except KeyError as e:
        print(f"欄位映射錯誤: {str(e)}")
        raise
    except Exception as e:
        print(f"創建欄位映射時發生未知錯誤: {str(e)}")
        raise

def format_items_table(doc, items_table, details, quote):
    """
    格式化項目表格
    
    參數:
    doc -- Document對象
    items_table -- 項目表格對象
    details -- 項目詳情列表
    quote -- 報價單數據
    """
    report_progress('processing', '正在處理項目表格', 40)
    
    # 保留標題行，刪除其他範例行
    if len(items_table.rows) > 1:
        # 刪除第一行以外的所有行
        for i in range(len(items_table.rows) - 1, 0, -1):
            try:
                items_table._element.remove(items_table.rows[i]._element)
            except Exception as e:
                print(f"刪除行時出錯: {e}")
    
    # 添加項目行
    total_items = len(details)
    for idx, item in enumerate(details):
        row = items_table.add_row()
        col_count = len(row.cells)
        
        # 設置每個單元格的值和格式
        if col_count >= 5:  # 標準五列表格 (類別, 項目, 單價, 數量, 金額)
            row.cells[0].text = item.get("category", "")
            row.cells[1].text = item.get("items", "")
            row.cells[2].text = format_number(item.get('unit', 0))
            row.cells[3].text = format_number(item.get('quantity', 0))
            row.cells[4].text = format_number(item.get('amount', 0))
            
            # 設置對齊方式
            apply_cell_style(row.cells[0], {"vertical_align": "center"})
            apply_cell_style(row.cells[1], {"vertical_align": "center"})
            apply_cell_style(row.cells[2], {"align": "right", "vertical_align": "center"})
            apply_cell_style(row.cells[3], {"align": "right", "vertical_align": "center"})
            apply_cell_style(row.cells[4], {"align": "right", "vertical_align": "center"})
            
        elif col_count == 4:  # 四列表格 (項目, 單價, 數量, 金額)
            row.cells[0].text = item.get("items", "")
            row.cells[1].text = format_number(item.get('unit', 0))
            row.cells[2].text = format_number(item.get('quantity', 0))
            row.cells[3].text = format_number(item.get('amount', 0))
            
            # 設置對齊方式
            apply_cell_style(row.cells[0], {"vertical_align": "center"})
            apply_cell_style(row.cells[1], {"align": "right", "vertical_align": "center"})
            apply_cell_style(row.cells[2], {"align": "right", "vertical_align": "center"})
            apply_cell_style(row.cells[3], {"align": "right", "vertical_align": "center"})
        
        print(f"已添加項目: {item.get('items', '')}")
        # 報告進度 (從40%開始，每個項目佔10%，到50%)
        progress = 40 + int((idx + 1) * 10 / total_items)
        report_progress('processing', f'處理項目 {idx+1}/{total_items}', progress)
    
    # 添加小計行
    report_progress('processing', '正在添加小計資訊', 55)
    subtotal_row = items_table.add_row()
    if len(subtotal_row.cells) >= 5:
        subtotal_row.cells[0].merge(subtotal_row.cells[3])
        subtotal_row.cells[0].text = "小計"
        subtotal_row.cells[4].text = format_number(quote.get('total_without_tax', 0))
        apply_cell_style(subtotal_row.cells[0], {"align": "right", "bold": True})
        apply_cell_style(subtotal_row.cells[4], {"align": "right", "bold": True})
    elif len(subtotal_row.cells) == 4:
        subtotal_row.cells[0].merge(subtotal_row.cells[2])
        subtotal_row.cells[0].text = "小計"
        subtotal_row.cells[3].text = format_number(quote.get('total_without_tax', 0))
        apply_cell_style(subtotal_row.cells[0], {"align": "right", "bold": True})
        apply_cell_style(subtotal_row.cells[3], {"align": "right", "bold": True})
    
    # 添加折扣行 (如果有)
    if quote.get("discount", 0) > 0:
        report_progress('processing', '正在添加折扣資訊', 60)
        discount_row = items_table.add_row()
        if len(discount_row.cells) >= 5:
            discount_row.cells[0].merge(discount_row.cells[3])
            discount_row.cells[0].text = "折扣"
            discount_row.cells[4].text = f"-{format_number(quote.get('discount', 0))}"
            apply_cell_style(discount_row.cells[0], {"align": "right"})
            apply_cell_style(discount_row.cells[4], {"align": "right"})
        elif len(discount_row.cells) == 4:
            discount_row.cells[0].merge(discount_row.cells[2])
            discount_row.cells[0].text = "折扣"
            discount_row.cells[3].text = f"-{format_number(quote.get('discount', 0))}"
            apply_cell_style(discount_row.cells[0], {"align": "right"})
            apply_cell_style(discount_row.cells[3], {"align": "right"})
    
    # 添加稅金行 (如果有)
    if quote.get("tax_rate", 0) > 0:
        report_progress('processing', '正在添加稅金資訊', 65)
        tax_row = items_table.add_row()
        if len(tax_row.cells) >= 5:
            tax_row.cells[0].merge(tax_row.cells[3])
            tax_row.cells[0].text = "稅金 (5%)"
            tax_row.cells[4].text = format_number(quote.get('tax_rate', 0))
            apply_cell_style(tax_row.cells[0], {"align": "right"})
            apply_cell_style(tax_row.cells[4], {"align": "right"})
        elif len(tax_row.cells) == 4:
            tax_row.cells[0].merge(tax_row.cells[2])
            tax_row.cells[0].text = "稅金 (5%)"
            tax_row.cells[3].text = format_number(quote.get('tax_rate', 0))
            apply_cell_style(tax_row.cells[0], {"align": "right"})
            apply_cell_style(tax_row.cells[3], {"align": "right"})
    
    # 添加總計行
    report_progress('processing', '正在添加總計資訊', 70)
    total_row = items_table.add_row()
    if len(total_row.cells) >= 5:
        total_row.cells[0].merge(total_row.cells[3])
        total_row.cells[0].text = "總計"
        total_row.cells[4].text = format_number(quote.get('total_with_tax', 0))
        apply_cell_style(total_row.cells[0], {"align": "right", "bold": True})
        apply_cell_style(total_row.cells[4], {"align": "right", "bold": True, "fill_color": "E6E6E6"})
    elif len(total_row.cells) == 4:
        total_row.cells[0].merge(total_row.cells[2])
        total_row.cells[0].text = "總計"
        total_row.cells[3].text = format_number(quote.get('total_with_tax', 0))
        apply_cell_style(total_row.cells[0], {"align": "right", "bold": True})
        apply_cell_style(total_row.cells[3], {"align": "right", "bold": True, "fill_color": "E6E6E6"})

def generate_docs(data):
    """
    生成報價單 Word 文檔
    
    參數:
    data -- 包含報價資訊的字典
    
    返回:
    list -- 生成的文檔本機路徑列表
    """
    try:
        # 轉換不同格式的輸入為標準格式
        standardized_data = standardize_input_data(data)
        
        # 驗證輸入數據格式
        if not isinstance(standardized_data, dict):
            raise ValueError(f"輸入數據必須是字典格式，實際類型: {type(standardized_data)}")
            
        if "quotes" not in standardized_data:
            raise ValueError("輸入數據缺少 'quotes' 字段")
            
        if not isinstance(standardized_data["quotes"], list) or len(standardized_data["quotes"]) == 0:
            raise ValueError("'quotes' 必須是非空列表")
            
        for idx, quote in enumerate(standardized_data["quotes"]):
            if not isinstance(quote, dict):
                raise ValueError(f"quote[{idx}] 必須是字典格式，實際類型: {type(quote)}")
                
            if "header" not in quote:
                raise ValueError(f"quote[{idx}] 缺少 'header' 字段")
                
            if "details" not in quote:
                raise ValueError(f"quote[{idx}] 缺少 'details' 字段")
        
        # 使用標準化後的數據生成文檔
        return generate_docs_from_template(standardized_data)
    except Exception as e:
        print(f"生成報價單時發生錯誤: {str(e)}")
        raise

def standardize_input_data(data):
    """
    將不同格式的輸入數據轉換為標準格式
    支持 Langflow 和直接 JSON 輸入
    
    參數:
    data -- 輸入數據，可能是各種格式
    
    返回:
    dict -- 標準化後的數據
    """
    print(f"原始輸入數據類型: {type(data)}")
    
    # 如果輸入為空，拋出錯誤
    if not data:
        print("輸入數據為空")
        raise ValueError("輸入數據為空")
    
    # 打印一下輸入數據的部分內容
    if isinstance(data, dict):
        print(f"輸入數據包含以下鍵: {list(data.keys())}")
        for key in data.keys():
            value_type = type(data[key])
            print(f"鍵 '{key}' 的值類型: {value_type}")
            if key == "quotes":
                # 處理特殊情況：quotes 是字典而不是列表（Langflow 可能的格式）
                if isinstance(data[key], dict):
                    print("檢測到 quotes 是字典而不是列表，將其轉換為標準格式")
                    # 將 quotes 值視為單個 quote
                    quote_dict = data[key]
                    # 檢查是否有 header 字段
                    if "header" in quote_dict:
                        print("quotes 字典中有 header 字段，將其視為單個 quote")
                        # 創建標準格式
                        data = {"quotes": [quote_dict]}
                        print(f"轉換後的數據: {data}")
                        return data
                    else:
                        print("quotes 字典中沒有 header 字段，創建缺省 header")
                        # 從字典中提取可能的 header 字段
                        header = {}
                        for possible_header in ["Title", "quoteNumber", "recipient", "companyName", "companyContact", "companyEmail", "start_date", "end_date", "key", "staff"]:
                            if possible_header in quote_dict:
                                header[possible_header] = quote_dict[possible_header]
                        
                        # 提取可能的 details 字段
                        details = []
                        if "details" in quote_dict:
                            details = quote_dict["details"]
                        elif "items" in quote_dict:
                            details = quote_dict["items"]
                        
                        # 創建標準格式
                        standardized_quote = {
                            "header": header,
                            "details": details
                        }
                        
                        # 復制其他字段
                        for k, v in quote_dict.items():
                            if k not in ["header", "details", "items"] + list(header.keys()):
                                standardized_quote[k] = v
                        
                        data = {"quotes": [standardized_quote]}
                        print(f"轉換後的數據: {data}")
                        return data
                
                elif isinstance(data[key], list):
                    print(f"quotes 列表長度: {len(data[key])}")
                    # 檢查列表中的每個項目是否為空
                    empty_quotes = True
                    for i, quote in enumerate(data[key]):
                        if isinstance(quote, dict):
                            print(f"quote[{i}] 包含的鍵: {list(quote.keys())}")
                            if quote and quote.keys():
                                empty_quotes = False
                    
                    # 如果所有 quote 都是空的，嘗試從本地文件加載預設數據
                    if empty_quotes:
                        print("檢測到所有 quotes 都是空的，嘗試加載備用數據")
                        backup_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "input.json")
                        if os.path.exists(backup_path):
                            try:
                                with open(backup_path, 'r', encoding='utf-8') as f:
                                    backup_data = json.load(f)
                                    print(f"已從 {backup_path} 載入備用數據")
                                    return backup_data
                            except Exception as e:
                                print(f"載入備用數據失敗: {str(e)}")
                        else:
                            print(f"未找到備用數據文件: {backup_path}")
    
    # 如果數據已經是標準格式，直接返回
    if isinstance(data, dict) and "quotes" in data and isinstance(data["quotes"], list):
        print("輸入數據已經是標準格式")
        # 檢查每個 quote 是否有 header 和 details
        for i, quote in enumerate(data["quotes"]):
            if not isinstance(quote, dict):
                print(f"quote[{i}] 不是字典類型")
                continue
                
            if "header" not in quote:
                print(f"quote[{i}] 缺少 header 字段，嘗試創建")
                # 從 quote 的頂層屬性構建 header
                quote["header"] = {}
                for key in list(quote.keys()):
                    if key not in ["details", "header", "total_without_tax", "discount", "tax_rate", "total_with_tax", "notes"]:
                        quote["header"][key] = quote.pop(key)
                print(f"為 quote[{i}] 創建的 header: {quote['header']}")
                
            if "details" not in quote:
                print(f"quote[{i}] 缺少 details 字段，設置為空列表")
                quote["details"] = []
        
        # 檢查數據是否有效 - 如果所有 quotes 的 header 都是空的
        all_empty_headers = True
        for quote in data["quotes"]:
            if quote.get("header") and quote["header"].keys():
                all_empty_headers = False
                break
                
        if all_empty_headers:
            print("所有 quotes 的 header 都是空的，嘗試載入備用數據")
            backup_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "input.json")
            if os.path.exists(backup_path):
                try:
                    with open(backup_path, 'r', encoding='utf-8') as f:
                        backup_data = json.load(f)
                        print(f"已從 {backup_path} 載入備用數據")
                        return backup_data
                except Exception as e:
                    print(f"載入備用數據失敗: {str(e)}")
        
        return data
    
    # 處理 Langflow 輸入格式：檢查是否為單個報價單數據（沒有外層 quotes 包裝）
    if isinstance(data, dict) and "header" not in data and "quotes" not in data:
        print("檢測到可能的非標準格式，嘗試從輸入構建標準格式")
        
        # 嘗試從輸入中構建標準格式
        # 情況1: 如果已有 Title, quoteNumber 等字段，將其視為 header
        header = {}
        details = []
        
        # 識別可能的 header 字段
        header_fields = [
            "Title", "quoteNumber", "recipient", "companyName", "companyContact",
            "companyEmail", "start_date", "end_date", "key", "staff"
        ]
        
        for field in header_fields:
            if field in data:
                header[field] = data[field]
                print(f"從輸入中找到 header 字段: {field} = {data[field]}")
        
        # 識別可能的 details 字段
        if "details" in data and isinstance(data["details"], list):
            details = data["details"]
            print(f"從輸入中找到 details 字段，包含 {len(details)} 項")
        elif "items" in data and isinstance(data["items"], list):
            details = data["items"]
            print(f"從輸入中找到 items 字段，包含 {len(details)} 項")
        
        # 其他可能的字段直接複製到 quote 根級別
        quote = {
            "header": header,
            "details": details
        }
        
        # 複製其他可能有用的字段
        for key, value in data.items():
            if key not in ["header", "details", "items"] + header_fields:
                quote[key] = value
                print(f"從輸入中複製其他字段: {key}")
        
        print(f"構建的標準格式 quote: {quote}")
        return {"quotes": [quote]}
    
    # 特殊情況：純文本或已經包含 header 字段，但沒有外層 quotes 包裝
    if isinstance(data, dict) and "header" in data:
        print("檢測到單個 quote 格式（含 header 但無 quotes 包裝）")
        return {"quotes": [data]}
    
    # 處理 Langflow 可能發送的其他格式
    try:
        # 嘗試以字符串形式解析 JSON
        if isinstance(data, str):
            print("輸入為字符串，嘗試解析 JSON")
            import json
            parsed_data = json.loads(data)
            print(f"解析後的數據類型: {type(parsed_data)}")
            return standardize_input_data(parsed_data)
    except Exception as e:
        print(f"解析 JSON 字符串失敗: {str(e)}")
    
    # 最後嘗試：將整個數據視為單個報價單
    try:
        print("嘗試將整個輸入視為單個報價單")
        
        standardized_data = {
            "quotes": [{
                "header": {
                    "Title": "報價單",
                    "quoteNumber": f"Q-{datetime.now().strftime('%Y-%m%d')}",
                    "recipient": "客戶",
                    "companyContact": "0988363357",
                    "companyEmail": "istudiodesign.tw@gmail.com",
                    "start_date": datetime.now().strftime("%Y/%m/%d"),
                    "end_date": (datetime.now().replace(month=datetime.now().month+1) if datetime.now().month < 12 else datetime.now().replace(year=datetime.now().year+1, month=1)).strftime("%Y/%m/%d"),
                    "key": "96790278",
                    "staff": "亦式數位互動有限公司"
                },
                "details": [],
                "total_without_tax": 0,
                "discount": 0,
                "tax_rate": 0,
                "total_with_tax": 0,
                "notes": "謝謝惠顧"
            }]
        }
        
        # 如果數據是字典，嘗試將其合併到第一個報價單中
        if isinstance(data, dict):
            print("將原始字典數據合併到標準格式中")
            for key, value in data.items():
                if key == "header" and isinstance(value, dict):
                    print(f"合併 header: {list(value.keys())}")
                    standardized_data["quotes"][0]["header"].update(value)
                elif key == "details" or key == "items":
                    if isinstance(value, list):
                        print(f"設置 details: {len(value)} 項")
                        standardized_data["quotes"][0]["details"] = value
                else:
                    print(f"複製字段: {key}")
                    standardized_data["quotes"][0][key] = value
        
        print(f"最終構建的標準格式: {standardized_data}")
        return standardized_data
    except Exception as e:
        print(f"無法將輸入數據轉換為標準格式: {str(e)}")
        raise ValueError(f"無法將輸入數據轉換為標準格式: {str(e)}")

def generate_docs_from_template(data):
    """
    使用模板生成報價單 Word 文檔
    
    參數:
    data -- 包含報價資訊的字典
    
    返回:
    list -- 生成的文檔本機路徑列表
    """
    outputs = []
    
    # 確保 temp 目錄存在
    report_progress('preparing', '準備處理環境', 0)
    temp_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "temp")
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir, exist_ok=True)
    
    # 清理舊檔案，避免權限衝突
    cleanup_temp_files(temp_dir)
    
    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "報價單.docx")
    
    if not os.path.exists(template_path):
        error_message = f"找不到模板檔案: {template_path}"
        report_progress('error', error_message, 0)
        raise FileNotFoundError(error_message)
    
    # 詳細分析模板結構
    template_info = analyze_template(template_path)
    
    total_quotes = len(data["quotes"])
    for idx, quote in enumerate(data["quotes"]):
        try:
            # 報告進度 - 每個報價單佔90%總進度的一部分
            progress_base = idx * 90 / total_quotes
            report_progress('processing', f'開始處理第 {idx+1}/{total_quotes} 份報價單', int(progress_base))
            
            # 載入模板
            doc = Document(template_path)
            
            # 確保必要字段存在
            if "header" not in quote:
                raise KeyError("報價單數據缺少 'header' 字段")
                
            if "details" not in quote:
                raise KeyError("報價單數據缺少 'details' 字段")
                
            header = quote["header"]
            details = quote.get("details", [])
            quote_number = header.get("quoteNumber", "unknown")
            print(f"生成報價單: {quote_number}")
            
            # 創建欄位映射
            report_progress('processing', '準備欄位映射', int(progress_base + 15))
            field_mapping = create_field_mapping(quote)
            
            print(f"\n欄位對應:")
            for key, value in field_mapping.items():
                print(f"  {key} → {value}")
            
            # 處理所有表格中的文字
            report_progress('processing', '處理表格佔位符', int(progress_base + 20))
            print("\n處理表格:")
            for t, table in enumerate(doc.tables):
                for r, row in enumerate(table.rows):
                    for c, cell in enumerate(row.cells):
                        for p, paragraph in enumerate(cell.paragraphs):
                            if "{" in paragraph.text and not ("{#items}" in paragraph.text or "{/items}" in paragraph.text):
                                print(f"表格 {t+1}, 行 {r+1}, 列 {c+1}, 段落 {p+1} 原始內容: '{paragraph.text}'")
                                if replace_text_with_field_value(paragraph, field_mapping):
                                    print(f"表格 {t+1}, 行 {r+1}, 列 {c+1}, 段落 {p+1} 已完成替換")
            
            # 尋找並填充項目表格
            if template_info["item_table_index"] >= 0 and template_info["item_table_index"] < len(doc.tables):
                items_table = doc.tables[template_info["item_table_index"]]
                print(f"處理項目表格 (索引 {template_info['item_table_index']})")
                
                # 格式化項目表格
                format_items_table(doc, items_table, details, quote)
            
            # 最後處理段落，跳過{#items}和{/items}標籤
            report_progress('processing', '處理文本佔位符', int(progress_base + 75))
            print("\n處理段落:")
            for i, paragraph in enumerate(doc.paragraphs):
                if "{" in paragraph.text and not ("{#items}" in paragraph.text or "{/items}" in paragraph.text):
                    print(f"段落 {i+1} 原始內容: '{paragraph.text}'")
                    if replace_text_with_field_value(paragraph, field_mapping):
                        print(f"段落 {i+1} 已完成替換")
            
            # 保存文件前嘗試刪除同名檔案
            report_progress('finalizing', '準備保存文檔', int(progress_base + 85))
            file_name = f"quote_{quote_number}.docx"
            file_path = os.path.join(temp_dir, file_name)
            
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                    print(f"已刪除舊檔案: {file_path}")
                    time.sleep(0.5)  # 等待一下以確保檔案被釋放
            except Exception as e:
                print(f"刪除舊檔案時出錯: {e}")
                # 使用時間戳來避免檔案名衝突
                file_name = f"quote_{quote_number}_{int(time.time())}.docx"
                file_path = os.path.join(temp_dir, file_name)
            
            # 保存文件
            doc.save(file_path)
            outputs.append(file_path)
            report_progress('finalizing', f'已生成報價單: {quote_number}', int(progress_base + 90))
            print(f"已成功生成報價單: {file_path}")
        except KeyError as e:
            error_message = f"處理報價單時發生欄位錯誤: {str(e)}"
            print(error_message)
            report_progress('error', error_message, 0)
        except Exception as e:
            error_message = f"處理報價單時發生錯誤: {str(e)}"
            print(error_message)
            report_progress('error', error_message, 0)
    
    report_progress('completed', f'已完成所有報價單處理, 共 {len(outputs)} 份', 100)
    return outputs

def main():
    # 讀取 JSON 資料檔
    try:
        with open('input.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 生成報價單
        doc_paths = generate_docs(data)
        
        print(f"\n已成功生成 {len(doc_paths)} 份報價單:")
        for path in doc_paths:
            print(f"- {path}")
    except Exception as e:
        print(f"程序執行時發生錯誤: {str(e)}")

if __name__ == "__main__":
    main() 